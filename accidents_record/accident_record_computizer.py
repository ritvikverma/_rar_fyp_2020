import json
import os

import docx

doc_files = []


def do_list_update():
    for file in os.listdir('./Acc_Record'):
        filename = os.fsdecode(file)
        if filename.endswith('docx') or filename.endswith('doc'):
            doc_files.append(os.path.join('./Acc_Record', filename))


def retrieve_log(file):
    doc = docx.Document(file)
    # Get totals
    total = dict(zip([cell.text for cell in doc.tables[0].rows[0].cells],
                     [cell.text for cell in doc.tables[0].rows[1].cells]))
    # Identify all events tables (blue one with event description)
    events = []
    for i in range(1, len(doc.tables)):
        current_table = doc.tables[i]
        if current_table.rows[0].cells[0].text.lower() == 'item':
            event = {}
            event_descriptions = []
            for row_index in range(1, len(current_table.rows)):
                current_table.rows[row_index].cells[0].text = current_table.rows[row_index].cells[0].text + \
                                                              '-' + \
                                                              str(row_index)
                event_descriptions.append(
                    dict(zip([cell.text.strip() for cell in current_table.rows[0].cells],
                             [cell.text.strip() for cell in current_table.rows[row_index].cells]))
                )
            event['event_descriptions'] = event_descriptions

            # Get event descriptions from next table
            next_table = doc.tables[i + 1]
            for row in next_table.rows:
                for cell in row.cells:
                    event[cell.text.split(':')[0].strip()] = cell.text.split(':')[1].strip()
            next_to_next_table = doc.tables[i + 2]

            # Get event descriptions from next to next table
            for row in next_to_next_table.rows:
                for cell in row.cells:
                    event[cell.text.split(':')[0].strip()] = cell.text.split(':')[1].strip()

            events.append(event)
    # Get traffic controllers
    controllers = {}
    footer_table = doc.sections[0].footer.tables[0]
    for i in range(1, len(footer_table.columns)):
        column = footer_table.columns[i]
        controllers[column.cells[0].text] = []
        for operator_cell in column.cells[1:]:
            controllers[column.cells[0].text].append(operator_cell.text)
    log = {
        'total': total,
        'events': events,
        'traffic_controllers': controllers
    }
    return log


def write_log(log, file_name):
    with open('./logs/' + file_name + '.json', 'w') as outfile:
        json.dump(log, outfile)


if __name__ == '__main__':
    do_list_update()
    for file in doc_files:
        file_name = file.split('/')[-1].split('.')[0]
        try:
            log = retrieve_log(file)
            write_log(log, file_name)
        except Exception as e:
            print(file_name, e)
